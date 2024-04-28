import lxml.etree as ET
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import parse_xml
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
import argparse
import os
import tracemalloc
import time
import xml.etree.ElementTree as ET

def parse_nessus_file(filename):
    severity_labels = {'0': 'Info', '1': 'Low', '2': 'Medium', '3': 'High', '4': 'Critical'}
    severity_order = {'Info': 0, 'Low': 1, 'Medium': 2, 'High': 3, 'Critical': 4}
    tree = ET.parse(filename)
    root = tree.getroot()
    all_vulnerabilities = {}
    host_vulnerabilities = []
    info_vulnerabilities = []
    
    for block in root.findall('.//ReportHost'):
        host_data = {"Hostname": block.get('name'), "Vulnerabilities": []}
        for vuln in block.findall('.//ReportItem'):
            severity_code = vuln.get('severity')
            severity = severity_labels.get(severity_code, 'Unknown')
            if severity == "Info":
                info_vuln_data = {
                    "Plugin Name": vuln.get('pluginName')
                }
                info_vulnerabilities.append(info_vuln_data)
                continue

            vuln_data = {
                "Plugin Name": vuln.get('pluginName'),
                "Severity": severity,
                "Synopsis": vuln.find('synopsis').text if vuln.find('synopsis') is not None else 'No synopsis provided',
                "Description": vuln.find('description').text if vuln.find('description') is not None else 'No description provided',
                "Solution": vuln.find('solution').text if vuln.find('solution') is not None and vuln.find('solution').text.strip() != '' else 'No remediation provided',
                "CVSS3 Base Score": vuln.find('cvss3_base_score').text if vuln.find('cvss3_base_score') is not None else 'N/A',
                "Exploitability Ease": vuln.find('exploitability_ease').text if vuln.find('exploitability_ease') is not None else 'Not available',
                "CVE": vuln.find('cvss_score_source').text if vuln.find('cvss_score_source') is not None else 'N/A'
            }
            host_data["Vulnerabilities"].append(vuln_data)
            if vuln_data["Plugin Name"] not in all_vulnerabilities or severity_order[all_vulnerabilities[vuln_data["Plugin Name"]]['Severity']] < severity_order[severity]:
                all_vulnerabilities[vuln_data["Plugin Name"]] = vuln_data
        host_vulnerabilities.append(host_data)
    return all_vulnerabilities, host_vulnerabilities, info_vulnerabilities

def create_pie_chart(all_vulns, filename):
    severities = {'Critical': 0, 'High': 0, 'Medium': 0, 'Low': 0}
    for vuln in all_vulns.values():
        severity = vuln['Severity']
        severities[severity] += 1

    labels = list(severities.keys())
    sizes = list(severities.values())

    fig1, ax1 = plt.subplots(figsize=(8, 6))  # Adjust the figure size here
    wedges, texts, autotexts = ax1.pie(sizes, labels=labels, autopct='%1.1f%%', startangle=90)

    # Add information about vulnerabilities to the pie chart
    vuln_info = [f"{label}: {size}" for label, size in zip(labels, sizes)]
    ax1.legend(wedges, vuln_info, title="Vulnerabilities", loc="center left", bbox_to_anchor=(1, 0, 0.5, 1))

    # Add a legend for the severity levels
    ax1.legend(loc='upper right', labels=severities.keys())

    ax1.axis('equal')

    plt.savefig(filename)
                                                                                                                                                                        


def apply_table_styles(table):
    table.style = 'Table Grid'
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Arial'
    for cell in table.rows[0].cells:
        if not cell.paragraphs[0].runs:
            run = cell.paragraphs[0].add_run()
        run.font.bold = True
        run.font.size = Pt(12)
        run.font.name = 'Arial'
        run.font.color.rgb = RGBColor(255, 255, 255)
        shading_elm = parse_xml(r'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="4F81BD"/>')
        cell._tc.get_or_add_tcPr().append(shading_elm)

def create_info_table(info_vulnerabilities, doc):
    if not info_vulnerabilities:
        return
    doc.add_heading('Executive Summary of Info Severities', level=1)
    table = doc.add_table(rows=1, cols=1)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Vulnerability'
    apply_table_styles(table)
    
    added_vulns = set()  # Keep track of added vulnerabilities
    for vuln in info_vulnerabilities:
        if vuln['Plugin Name'] not in added_vulns:
            row_cells = table.add_row().cells
            row_cells[0].text = vuln['Plugin Name']
            added_vulns.add(vuln['Plugin Name'])

def create_docx(all_vulns, host_vulns, info_vulns, docx_filename, option):
    doc = Document()
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(11)

    if option in ['1', '2']:  # Overview and Full Report
        doc.add_heading('Summary of Vulnerabilities', level=1)
        # Create the summary table
        summary_table = doc.add_table(rows=1, cols=4)
        summary_hdr_cells = summary_table.rows[0].cells
        summary_hdr_cells[0].text = 'Vulnerability'
        summary_hdr_cells[1].text = 'Severity'
        summary_hdr_cells[2].text = 'CVSS3 Base Score'
        summary_hdr_cells[3].text = 'CVE'
        apply_table_styles(summary_table)
        for details in all_vulns.values():
            row_cells = summary_table.add_row().cells
            row_cells[0].text = details["Plugin Name"]
            row_cells[1].text = details["Severity"]
            row_cells[2].text = details["CVSS3 Base Score"]
            row_cells[3].text = details.get("CVE", "N/A")

        # Generate and include the pie chart
        pie_chart_path = f"{os.path.splitext(docx_filename)[0]}_pie_chart.png"
        create_pie_chart(all_vulns, pie_chart_path)
        doc.add_picture(pie_chart_path, width=Inches(4))
        os.remove(pie_chart_path)

        # Add the remediation table
        doc.add_heading('Remediation per Vulnerability', level=1)
        remediation_table = doc.add_table(rows=1, cols=2)
        remediation_hdr_cells = remediation_table.rows[0].cells
        remediation_hdr_cells[0].text = 'Vulnerability'
        remediation_hdr_cells[1].text = 'Recommended Action'
        apply_table_styles(remediation_table)
        for details in all_vulns.values():
            row_cells = remediation_table.add_row().cells
            row_cells[0].text = details["Plugin Name"]
            row_cells[1].text = details["Solution"]

    if option in ['1', '3']:  # Vulnerability per Host and Full Report
        for host in host_vulns:
            doc.add_heading(f'Host: {host["Hostname"]}', level=2)
            table = doc.add_table(rows=1, cols=6)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Severity'
            hdr_cells[1].text = 'Plugin Name'
            hdr_cells[2].text = 'Synopsis'
            hdr_cells[3].text = 'Description'
            hdr_cells[4].text = 'Solution'
            hdr_cells[5].text = 'Exploitability Ease'
            apply_table_styles(table)
            for vuln in host['Vulnerabilities']:
                row_cells = table.add_row().cells
                row_cells[0].text = vuln['Severity']
                row_cells[1].text = vuln['Plugin Name']
                row_cells[2].text = vuln['Synopsis']
                row_cells[3].text = vuln['Description']
                row_cells[4].text = vuln['Solution']
                row_cells[5].text = vuln['Exploitability Ease']

    if option in ['1', '2']:  # Info severity executive summary
        create_info_table(info_vulns, doc)

    doc.save(docx_filename)
class MonitorPerformance:
    def __enter__(self):
        tracemalloc.start()
        self.start_time = time.time()
        return self

    def __exit__(self, exc_type, exc_val, exc_tb):
        self.elapsed_time = time.time() - self.start_time
        self.current, self.peak = tracemalloc.get_traced_memory()
        tracemalloc.stop()
        print(f"Elapsed Time: {self.elapsed_time:.2f} seconds")
        print(f"Current memory usage is {self.current / 10**6:.2f} MB; Peak was {self.peak / 10**6:.2f} MB")

# Usage in the main function:
def main():
    parser = argparse.ArgumentParser(description="Parse Nessus XML files and generate reports.")
    parser.add_argument("file", help="The Nessus XML file to parse.")
    parser.add_argument("option", choices=['1', '2', '3', '4'], help="1: Full report, 2: Overview, 3: Per host, 4: Remediations")
    args = parser.parse_args()

    with MonitorPerformance():
        all_vulnerabilities, host_vulnerabilities, info_vulnerabilities = parse_nessus_file(args.file)
        output_docx = f"Nessus_Report_{args.option}.docx"
        create_docx(all_vulnerabilities, host_vulnerabilities, info_vulnerabilities, output_docx, args.option)
        print(f'Report generated: {output_docx}')

if __name__ == "__main__":
    main()

